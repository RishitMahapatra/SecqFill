"""
The main purpose of this code is to parse the excel sheets where the questionnaires are mentioned

The part to be handled is that most of the companies have their own way of demonstrating the questionnaries in excel sheet 

The hard part is that every enterprise buyer's spreadsheet has a different
layout: different column order, different headers, sometimes no header at
all. So instead of hardcoding "questions are in column B," we score every
column on how "question-like" its contents are, and pick the winner.
"""
import re #imports request 
from dataclasses import dataclass 
import openpyxl
from docx import Document 

# Common sentence openers in security/compliance questions. This is a cheap
# signal, not a real NLP classifier — regex is enough because we only need
# it to nudge the score in ambiguous cases, not to understand meaning.
QUESTION_STEMS = re.compile(
    r"^(do you|does the|is there|are there|what is|where is|how do|"
    r"can you|have you|will you|who is|when was)",
    re.IGNORECASE,
)

@dataclass #the purpose of dataclass is to store data and state with minimal boiler plate or blue print 
class ExtractedQuestion:        #This means ExtractedQuestion is a boiler plate class for dataclass 
    row_number: int             #The row number in the excel sheet 
    question_text : str         #This is the plain text of the question


def _column_score(cell_values: list[str]) -> float : #this provides a score to how much likely the current column contains the question
    #the logic for the same is provided below 
    """
    Score ONE column on how likely it is to be the question column.
    Higher score = more likely. Shared by both the xlsx and docx parsers
    below, because a spreadsheet column and a Word table column are
    structurally the same thing: a list of text values.

    The reasoning behind each signal:

    1. avg_length — the single strongest signal. A "Question" or
       "Requirement" column is made of full sentences. An "Answer" or
       "Notes" column next to it is usually empty, or short things like
       "Yes," "N/A," "See appendix." Long average text = probably questions.

    2. question_mark_hits — a direct, almost-zero-false-positive signal.
       If a cell ends in "?", it's overwhelmingly likely to be a question,
       not an answer or a label. Weighted heavily (x15) because it's rare
       to see a "?" anywhere except in an actual question.

    3. stem_hits — a secondary signal for questions that are phrased as
       statements rather than literal questions ("Provide details of your
       incident response plan" has no "?" but is clearly the question
       column). Weighted lower (x10) than the "?" signal because phrasing
       varies more and this is a noisier match.
    """
    non_empty = [v for v in cell_values if v and v.strip()] #checks and stores the values for cell values containing some characters of interest 
    # this filters out the missing cells or only white spaces values with the help of strip keyword 

    if not non_empty: #if the cell value does not exist in the list of non empty values then the col is not containing anything imp so socre is 0.0
        return 0.0 

    #calculating the avg length of the valid text cells 
    avg_length = sum(len(v) for v in non_empty) / len(non_empty)
    #total sum of the lengths of the cell values by the non empty list length to get the average 

    #calculating the total number of ? in the plain text of the current cell value
    question_mark_hits = sum (1 for v in non_empty if "?" in v)
    stem_hits = sum(1 for v in non_empty if QUESTION_STEMS.match(v.strip()))
    
    score = avg_length
    score += question_mark_hits * 15
    score += stem_hits * 10
    return score

def extract_questions_from_xlsx(file_path: str) -> list[ExtractedQuestion]:
    # data_only=True is important: if a cell contains a formula, this
    # returns the last-calculated VALUE instead of the formula string
    # itself. We want "Yes" not "=IF(A1=1,"Yes","No")".
    wb = openpyxl.load_workbook(file_path, data_only=True) #opens and reads the excel sheet 

    #AS OF NOW ASSUMES MOST OF THE QUESTIONS WILL BE WITHIN A SINGLE TAB 
    ws = wb.active  #the active first page/tab 
    all_rows = list(ws.iter_rows(values_only=True)) #only the text info needed rather than the formatting information 

    if not all_rows:
        return []
    
    num_cols = max(len(r) for r in all_rows)
    columns: list[list[str]] = [[] for _ in range(num_cols)]
    for row in all_rows:
        for col_idx in range(num_cols):
            value = row[col_idx] if col_idx < len(row) else None
            # Cast everything to str now — Excel cells can hold numbers,
            # dates, booleans, None. Normalizing to strings early means
            # every downstream function only has one type to worry about.
            columns[col_idx].append(str(value) if value is not None else "")

    scores = [_column_score(col) for col in columns]
    # argmax over the scores — "which column index scored highest"
    question_col_idx = max(range(num_cols), key=lambda i: scores[i])
    results = []
    for row_num, row in enumerate(all_rows, start=1):
        value = row[question_col_idx] if question_col_idx < len(row) else None
        text = str(value).strip() if value is not None else ""

        if not text:
            continue  # empty cell, nothing to extract

        # Deliberately narrow header check: only trigger on row 1, only if
        # short, and only against a small known list of header words.
        # A real question could theoretically be short and coincidentally
        # match one of these words — being narrow here trades a small
        # false-negative risk (rarely misses a header) for avoiding the
        # bigger risk of accidentally discarding a real short question.
        if len(text) < 30 and text.lower() in (  #ignores the basic keywords
            "question", "requirement", "control", "item", "questions",
        ):
            continue

        results.append(ExtractedQuestion(row_number=row_num, question_text=text))

    return results

def extract_questions_from_docx(file_path: str) -> list[ExtractedQuestion]:
    doc = Document(file_path)
    if not doc.tables:
        return [] #This means if there are no table in the document then return nothing as this function can extract nothing 
    
    #in case of the table exists 
    table = doc.tables[0]
    # python-docx gives us table.rows -> each row has .cells -> each cell
    # has .text. This list comprehension flattens that into the exact
    # same "list of rows, each a list of strings" shape the xlsx version
    # produces — which is what lets us reuse the transpose/score logic
    # below almost verbatim.
    all_rows = [[cell.text for cell in row.cells] for row in table.rows]
    if not all_rows:
        return []
    num_cols = max(len(r) for r in all_rows)
    columns: list[list[str]] = [[] for _ in range(num_cols)]
    for row in all_rows:
        for col_idx in range(num_cols):
            columns[col_idx].append(row[col_idx] if col_idx < len(row) else "")

    scores = [_column_score(col) for col in columns]
    question_col_idx = max(range(num_cols), key=lambda i: scores[i])
    results = []
    for row_num, row in enumerate(all_rows, start=1):
        text = row[question_col_idx].strip() if question_col_idx < len(row) else ""
        if not text:
            continue
        if len(text) < 30 and text.lower() in ( #ignores the basic keywords
            "question", "requirement", "control", "item", "questions",
        ):
            continue
        results.append(ExtractedQuestion(row_number=row_num, question_text=text))

    return results

    


